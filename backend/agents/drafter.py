import os
import sys
import json
import requests
import uuid
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from groq import Groq
from docx import Document
from dotenv import load_dotenv
from database import SessionLocal, GapReport, CompanyPolicy, AuditTrail
from datetime import datetime

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

SLACK_WEBHOOK_URL = os.getenv("SLACK_WEBHOOK_URL")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")
JIRA_EMAIL = os.getenv("JIRA_EMAIL")
JIRA_URL = os.getenv("JIRA_URL")
JIRA_PROJECT_KEY = os.getenv("JIRA_PROJECT_KEY")

def log_audit(pipeline_run_id, action, decision, branch=None, confidence=None, regulation_title=None):
    db = SessionLocal()
    try:
        db.add(AuditTrail(
            pipeline_run_id=pipeline_run_id,
            agent_name="Drafter",
            action=action,
            decision=decision,
            branch_taken=branch or "",
            confidence=confidence,
            regulation_title=regulation_title or ""
        ))
        db.commit()
    finally:
        db.close()

def save_gaps_to_db(gap_data: dict, jira_key: str, doc_path: str):
    db = SessionLocal()
    try:
        for gap in gap_data.get("gaps", []):
            db.add(GapReport(
                regulation_title=gap_data.get("regulation_title", ""),
                regulation_hash=gap_data.get("regulation_hash", ""),
                gap_description=gap.get("gap", ""),
                confidence_score=gap.get("confidence_score", 0.0),
                confidence_reason=gap.get("confidence_reason", ""),
                policy_section=gap.get("policy_section", ""),
                jurisdiction=gap_data.get("jurisdiction", ""),
                deadline=gap_data.get("deadline", ""),
                days_remaining=gap_data.get("days_remaining", -1),
                enforcement_context=gap_data.get("enforcement_context", ""),
                jira_key=jira_key,
                doc_path=doc_path,
                requires_human_review=gap.get("confidence_score", 1.0) < 0.7
            ))
        db.commit()
        print(f"  ✅ {len(gap_data.get('gaps', []))} gaps saved to database")
    finally:
        db.close()

def update_policy_in_db(gap: dict, suggested_content: str):
    db = SessionLocal()
    try:
        section = gap.get("policy_section", "")
        policy = db.query(CompanyPolicy).filter(
            CompanyPolicy.section.contains(section.split("—")[0].strip())
        ).first()
        if policy:
            policy.content = suggested_content
            policy.last_updated = datetime.utcnow()
            policy.updated_by = "ComplianceBot Agent"
            old_version = str(policy.version)
            parts = old_version.split(".")
            if len(parts) >= 2:
                policy.version = parts[0] + "." + str(int(parts[1]) + 1)
            else:
                policy.version = parts[0] + ".1"
            print(f"  📝 Policy updated in DB: {policy.title} → v{policy.version}")
    finally:
        db.close()

def create_jira_ticket(regulation_title, gaps, deadline, days_remaining, enforcement_context):
    """USB 07 — Deadline countdown in Jira ticket"""
    if not all([JIRA_URL, JIRA_EMAIL, JIRA_API_TOKEN, JIRA_PROJECT_KEY]):
        print("  ⚠️  Jira not configured")
        return None

    deadline_text = f"\n📅 *Compliance Deadline:* {deadline} ({days_remaining} days remaining)" if deadline else ""
    enforcement_text = f"\n\n🚨 *Recent Enforcement Actions:*\n{enforcement_context}" if enforcement_context else ""

    gaps_text = "\n".join([
        f"• [{g.get('confidence_score', 0):.0%} confidence] {g.get('gap', '')} — Fix: {g.get('suggested_fix', '')}"
        for g in gaps
    ])

    description_text = (
        f"Regulation: {regulation_title}"
        f"{deadline_text}"
        f"\n\nGaps Identified:\n{gaps_text}"
        f"{enforcement_text}"
    )

    jira_data = {
        "fields": {
            "project": {"key": JIRA_PROJECT_KEY},
            "summary": f"Compliance Gap: {regulation_title[:80]}",
            "description": {
                "type": "doc", "version": 1,
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": description_text}]}]
            },
            "issuetype": {"name": "Task"}
        }
    }

    try:
        res = requests.post(
            f"{JIRA_URL}/rest/api/3/issue",
            json=jira_data,
            headers={"Accept": "application/json", "Content-Type": "application/json"},
            auth=(JIRA_EMAIL, JIRA_API_TOKEN)
        )
        if res.status_code in [200, 201]:
            key = res.json().get("key")
            print(f"  🎫 Jira ticket created: {key}{deadline_text}")
            return key
    except Exception as e:
        print(f"  ❌ Jira error: {e}")
    return None

def send_slack(regulation_title, gaps_count, jira_key, deadline, days_remaining, requires_human_review):
    if not SLACK_WEBHOOK_URL:
        return
    deadline_text = f"\n📅 Deadline: {deadline} (*{days_remaining} days remaining*)" if deadline else ""
    review_text = "\n⚠️ *Human review required* — confidence below threshold" if requires_human_review else ""
    msg = {
        "text": (
            f"🤖 *ComplianceBot Alert*\n"
            f"*Regulation:* {regulation_title}\n"
            f"*Gaps Found:* {gaps_count}"
            f"{deadline_text}"
            f"\n*Jira:* {jira_key or 'N/A'}"
            f"{review_text}"
        )
    }
    try:
        requests.post(SLACK_WEBHOOK_URL, json=msg)
        print("  💬 Slack message sent")
    except Exception as e:
        print(f"  ❌ Slack error: {e}")

async def run_drafter(gap_data: dict, pipeline_run_id: str = None):
    if not pipeline_run_id:
        pipeline_run_id = gap_data.get("pipeline_run_id", str(uuid.uuid4())[:8])

    print("✍️  Agent 4: Drafter creating outputs...")

    regulation_title = gap_data.get("regulation_title", "Unknown Regulation")
    gaps = gap_data.get("gaps", [])
    deadline = gap_data.get("deadline", "")
    days_remaining = gap_data.get("days_remaining", -1)
    enforcement_context = gap_data.get("enforcement_context", "")
    jurisdiction = gap_data.get("jurisdiction", "US")

    requires_human_review = any(g.get("confidence_score", 1.0) < 0.7 for g in gaps)

    if not gaps:
        print("  ℹ️  No gaps to draft")
        return {"gaps_count": 0, "jira_key": None, "doc_path": None}

    # Generate SOP updates with Groq
    gaps_summary = "\n".join([f"- {g['gap']} (Fix: {g.get('suggested_fix', '')})" for g in gaps])
    prompt = f"""Write updated SOP policy sections for these compliance gaps:

Regulation: {regulation_title}
Jurisdiction: {jurisdiction}
{f'Deadline: {deadline} ({days_remaining} days remaining)' if deadline else ''}

Gaps:
{gaps_summary}

Write 2-3 paragraphs of professional SOP policy language that addresses each gap.
Be specific, use compliance terminology."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "You are a compliance policy writer."},
            {"role": "user", "content": prompt}
        ]
    )
    sop_text = response.choices[0].message.content

    # Save Word doc
    os.makedirs("outputs", exist_ok=True)
    safe_title = "".join(c for c in regulation_title if c.isalnum() or c in " _-")[:40]
    doc_path = f"outputs/SOP_Update_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    doc = Document()
    doc.add_heading("ComplianceBot — SOP Update", 0)
    doc.add_heading(regulation_title, level=1)

    if deadline:
        doc.add_paragraph(f"⚠️ Compliance Deadline: {deadline} ({days_remaining} days remaining)")

    doc.add_heading("Gaps Identified", level=2)
    for g in gaps:
        score = g.get("confidence_score", 0)
        doc.add_paragraph(
            f"[{score:.0%} confidence] {g.get('gap', '')}\n"
            f"Reason: {g.get('confidence_reason', '')}\n"
            f"Suggested Fix: {g.get('suggested_fix', '')}",
            style="List Bullet"
        )

    doc.add_heading("Updated SOP Language", level=2)
    doc.add_paragraph(sop_text)

    if enforcement_context and enforcement_context != "No recent enforcement actions found.":
        doc.add_heading("Recent Enforcement Actions", level=2)
        doc.add_paragraph(enforcement_context)

    doc.add_heading("Audit Information", level=2)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\nPipeline Run: {pipeline_run_id}\nJurisdiction: {jurisdiction}")

    doc.save(doc_path)
    print(f"  📄 Word doc saved: {doc_path}")

    # Update policy in DB
    for gap in gaps:
        update_policy_in_db(gap, sop_text)

    # Create Jira ticket
    jira_key = create_jira_ticket(regulation_title, gaps, deadline, days_remaining, enforcement_context)

    # Send Slack
    send_slack(regulation_title, len(gaps), jira_key, deadline, days_remaining, requires_human_review)

    # Save gap reports to DB
    gap_data["enforcement_context"] = enforcement_context
    save_gaps_to_db(gap_data, jira_key or "", doc_path)

    log_audit(pipeline_run_id, "drafted_outputs", f"Doc: {doc_path}, Jira: {jira_key}", regulation_title=regulation_title)

    return {
        "gaps_count": len(gaps),
        "jira_key": jira_key,
        "doc_path": doc_path,
        "requires_human_review": requires_human_review
    }